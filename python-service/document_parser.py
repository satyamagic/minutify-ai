import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict, List, Any
import re


def parse_pdf(pdf_path: str) -> Dict[str, Any]:
    """
    Parse PDF file and extract text in logical sections
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Dict containing segments organized by section
    """
    try:
        segments = []
        
        # Use pdfplumber for better text extraction
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            current_section = []
            current_title = None
            
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                
                if not text:
                    continue
                
                # Split into lines
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    
                    if not line:
                        continue
                    
                    # Detect potential headings (simple heuristic)
                    # Headings are typically: short, all caps, or end with colon
                    is_heading = (
                        len(line) < 100 and
                        (line.isupper() or 
                         re.match(r'^[\d\.]+\s+[A-Z]', line) or
                         line.endswith(':'))
                    )
                    
                    if is_heading and current_section:
                        # Save previous section
                        segments.append({
                            "index": len(segments),
                            "minute": None,
                            "title": current_title,
                            "content": "\n".join(current_section)
                        })
                        current_section = []
                        current_title = line
                    elif is_heading:
                        current_title = line
                    else:
                        current_section.append(line)
            
            # Add remaining content
            if current_section:
                segments.append({
                    "index": len(segments),
                    "minute": None,
                    "title": current_title,
                    "content": "\n".join(current_section)
                })
            
            # If no sections were detected, create one big section
            if not segments:
                full_text = "\n".join([
                    page.extract_text() for page in pdf.pages
                    if page.extract_text()
                ])
                segments.append({
                    "index": 0,
                    "minute": None,
                    "title": "Document Content",
                    "content": full_text
                })
        
        return {
            "segments": segments,
            "pages": total_pages
        }
        
    except Exception as e:
        raise Exception(f"PDF parsing failed: {str(e)}")


def parse_docx(docx_path: str) -> Dict[str, Any]:
    """
    Parse DOCX file and extract text in logical sections
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dict containing segments organized by section
    """
    try:
        doc = Document(docx_path)
        segments = []
        current_section = []
        current_title = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_section:
                    segments.append({
                        "index": len(segments),
                        "minute": None,
                        "title": current_title,
                        "content": "\n".join(current_section)
                    })
                    current_section = []
                
                current_title = text
            else:
                current_section.append(text)
        
        # Add remaining content
        if current_section:
            segments.append({
                "index": len(segments),
                "minute": None,
                "title": current_title,
                "content": "\n".join(current_section)
            })
        
        # If no sections were detected, create one big section
        if not segments:
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            segments.append({
                "index": 0,
                "minute": None,
                "title": "Document Content",
                "content": full_text
            })
        
        return {
            "segments": segments
        }
        
    except Exception as e:
        raise Exception(f"DOCX parsing failed: {str(e)}")
